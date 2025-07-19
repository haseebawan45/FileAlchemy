import os
import gc
import tempfile
from typing import Dict, Callable, Optional
import logging
from resource_manager import get_resource_manager

logger = logging.getLogger(__name__)

class OptimizedConverter:
    """Memory-efficient conversion methods optimized for Railway's free tier"""
    
    def __init__(self):
        self.resource_manager = get_resource_manager()
        
        # Lightweight conversion methods prioritized
        self.conversion_methods = {
            'pdf': {
                'text': self.lightweight_pdf_to_text,
                'docx': self.streaming_pdf_to_docx,
                'html': self.lightweight_pdf_to_html,
            },
            'docx': {
                'pdf': self.lightweight_docx_to_pdf,
                'txt': self.lightweight_docx_to_text,
                'html': self.lightweight_docx_to_html,
            },
            'image': {
                'convert': self.streaming_image_convert,
                'pdf': self.lightweight_image_to_pdf,
                'ocr': self.memory_efficient_ocr,
            },
            'xlsx': {
                'csv': self.lightweight_xlsx_to_csv,
                'json': self.streaming_xlsx_to_json,
            }
        }
    
    def select_conversion_method(self, file_type: str, target_format: str) -> Optional[Callable]:
        """Select the most memory-efficient conversion method"""
        if file_type in self.conversion_methods:
            if target_format in self.conversion_methods[file_type]:
                return self.conversion_methods[file_type][target_format]
        return None
    
    def lightweight_pdf_to_text(self, pdf_path: str, task_id: str = None) -> str:
        """Memory-efficient PDF to text conversion"""
        try:
            import fitz  # PyMuPDF
            
            if task_id:
                from server import conversion_progress
                conversion_progress[task_id] = {"progress": 20, "status": "Opening PDF"}
            
            txt_path = os.path.join("converted", os.path.basename(pdf_path).replace(".pdf", "_converted.txt"))
            
            # Process PDF page by page to minimize memory usage
            with fitz.open(pdf_path) as pdf_document:
                total_pages = len(pdf_document)
                
                with open(txt_path, "w", encoding="utf-8") as txt_file:
                    for page_num in range(total_pages):
                        if task_id:
                            progress = 20 + (page_num / total_pages) * 70
                            conversion_progress[task_id] = {
                                "progress": int(progress), 
                                "status": f"Processing page {page_num + 1}/{total_pages}"
                            }
                        
                        page = pdf_document.load_page(page_num)
                        text = page.get_text()
                        txt_file.write(f"--- Page {page_num + 1} ---\n")
                        txt_file.write(text)
                        txt_file.write("\n\n")
                        
                        # Force cleanup after each page
                        del page
                        if page_num % 5 == 0:  # Cleanup every 5 pages
                            gc.collect()
            
            if task_id:
                conversion_progress[task_id] = {"progress": 100, "status": "Conversion complete"}
            
            return txt_path
            
        except Exception as e:
            logger.error(f"PDF to text conversion failed: {e}")
            raise
    
    def streaming_pdf_to_docx(self, pdf_path: str, task_id: str = None) -> str:
        """Streaming PDF to DOCX conversion with memory management"""
        try:
            import fitz
            from docx import Document
            
            if task_id:
                from server import conversion_progress
                conversion_progress[task_id] = {"progress": 20, "status": "Starting PDF to DOCX conversion"}
            
            docx_path = os.path.join("converted", os.path.basename(pdf_path).replace(".pdf", "_converted.docx"))
            
            doc = Document()
            doc.add_heading(os.path.basename(pdf_path).replace(".pdf", ""), 0)
            
            with fitz.open(pdf_path) as pdf_document:
                total_pages = len(pdf_document)
                
                for page_num in range(total_pages):
                    if task_id:
                        progress = 20 + (page_num / total_pages) * 70
                        conversion_progress[task_id] = {
                            "progress": int(progress),
                            "status": f"Converting page {page_num + 1}/{total_pages}"
                        }
                    
                    page = pdf_document.load_page(page_num)
                    
                    # Add page heading
                    if page_num > 0:
                        doc.add_heading(f"Page {page_num + 1}", level=1)
                    
                    # Extract text in blocks to preserve some formatting
                    blocks = page.get_text("dict")["blocks"]
                    for block in blocks:
                        if block["type"] == 0:  # Text block
                            if "lines" in block:
                                for line in block["lines"]:
                                    if "spans" in line:
                                        line_text = ""
                                        for span in line["spans"]:
                                            line_text += span.get("text", "")
                                        if line_text.strip():
                                            doc.add_paragraph(line_text.strip())
                    
                    # Cleanup after each page
                    del page, blocks
                    if page_num % 3 == 0:  # More frequent cleanup for DOCX
                        gc.collect()
            
            doc.save(docx_path)
            
            if task_id:
                conversion_progress[task_id] = {"progress": 100, "status": "DOCX conversion complete"}
            
            return docx_path
            
        except Exception as e:
            logger.error(f"PDF to DOCX conversion failed: {e}")
            raise
    
    def lightweight_pdf_to_html(self, pdf_path: str, task_id: str = None) -> str:
        """Lightweight PDF to HTML conversion"""
        try:
            import fitz
            
            if task_id:
                from server import conversion_progress
                conversion_progress[task_id] = {"progress": 20, "status": "Starting PDF to HTML conversion"}
            
            html_path = os.path.join("converted", os.path.basename(pdf_path).replace(".pdf", "_converted.html"))
            
            html_content = """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        .page {{ margin-bottom: 40px; page-break-after: always; }}
        .page-header {{ font-weight: bold; color: #666; margin-bottom: 20px; }}
    </style>
</head>
<body>
""".format(title=os.path.basename(pdf_path))
            
            with fitz.open(pdf_path) as pdf_document:
                total_pages = len(pdf_document)
                
                for page_num in range(total_pages):
                    if task_id:
                        progress = 20 + (page_num / total_pages) * 70
                        conversion_progress[task_id] = {
                            "progress": int(progress),
                            "status": f"Converting page {page_num + 1}/{total_pages}"
                        }
                    
                    page = pdf_document.load_page(page_num)
                    text = page.get_text()
                    
                    html_content += f'<div class="page">\n'
                    html_content += f'<div class="page-header">Page {page_num + 1}</div>\n'
                    
                    # Simple text to HTML conversion
                    paragraphs = text.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            # Escape HTML characters
                            para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            html_content += f'<p>{para.strip()}</p>\n'
                    
                    html_content += '</div>\n'
                    
                    # Cleanup
                    del page
                    if page_num % 5 == 0:
                        gc.collect()
            
            html_content += """
</body>
</html>"""
            
            with open(html_path, "w", encoding="utf-8") as html_file:
                html_file.write(html_content)
            
            if task_id:
                conversion_progress[task_id] = {"progress": 100, "status": "HTML conversion complete"}
            
            return html_path
            
        except Exception as e:
            logger.error(f"PDF to HTML conversion failed: {e}")
            raise
    
    def streaming_image_convert(self, image_path: str, target_format: str, task_id: str = None) -> str:
        """Memory-efficient image conversion with streaming"""
        try:
            from PIL import Image
            
            if task_id:
                from server import conversion_progress
                conversion_progress[task_id] = {"progress": 20, "status": f"Converting to {target_format.upper()}"}
            
            # Determine output path
            base_name = os.path.splitext(os.path.basename(image_path))[0]
            output_path = os.path.join("converted", f"{base_name}_converted.{target_format.lower()}")
            
            # Open and convert image with memory optimization
            with Image.open(image_path) as img:
                # Convert mode if necessary
                if target_format.upper() == 'JPEG' and img.mode in ('RGBA', 'LA', 'P'):
                    # Convert to RGB for JPEG
                    rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    rgb_img.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                    img = rgb_img
                
                if task_id:
                    conversion_progress[task_id] = {"progress": 70, "status": "Saving converted image"}
                
                # Save with optimization
                save_kwargs = {'optimize': True}
                if target_format.upper() == 'JPEG':
                    save_kwargs['quality'] = 85
                elif target_format.upper() == 'PNG':
                    save_kwargs['compress_level'] = 6
                
                img.save(output_path, format=target_format.upper(), **save_kwargs)
            
            if task_id:
                conversion_progress[task_id] = {"progress": 100, "status": "Image conversion complete"}
            
            return output_path
            
        except Exception as e:
            logger.error(f"Image conversion failed: {e}")
            raise
    
    def memory_efficient_ocr(self, image_path: str, task_id: str = None) -> str:
        """Memory-efficient OCR processing"""
        try:
            import easyocr
            from PIL import Image
            
            if task_id:
                from server import conversion_progress
                conversion_progress[task_id] = {"progress": 20, "status": "Initializing OCR"}
            
            txt_path = os.path.join("converted", os.path.basename(image_path).replace(
                os.path.splitext(image_path)[1], "_ocr.txt"))
            
            # Initialize OCR reader (reuse global instance if available)
            try:
                from server import ocr_reader
                reader = ocr_reader
            except:
                reader = easyocr.Reader(['en'])
            
            if task_id:
                conversion_progress[task_id] = {"progress": 50, "status": "Processing image with OCR"}
            
            # Process image
            results = reader.readtext(image_path)
            
            if task_id:
                conversion_progress[task_id] = {"progress": 80, "status": "Extracting text"}
            
            # Extract text
            extracted_text = ""
            for (bbox, text, confidence) in results:
                if confidence > 0.5:  # Only include high-confidence text
                    extracted_text += text + "\n"
            
            # Save extracted text
            with open(txt_path, "w", encoding="utf-8") as txt_file:
                txt_file.write(f"OCR Results from: {os.path.basename(image_path)}\n")
                txt_file.write("=" * 50 + "\n\n")
                txt_file.write(extracted_text)
            
            if task_id:
                conversion_progress[task_id] = {"progress": 100, "status": "OCR complete"}
            
            # Force cleanup
            gc.collect()
            
            return txt_path
            
        except Exception as e:
            logger.error(f"OCR processing failed: {e}")
            raise
    
    def lightweight_xlsx_to_csv(self, xlsx_path: str, task_id: str = None) -> str:
        """Memory-efficient Excel to CSV conversion"""
        try:
            import pandas as pd
            
            if task_id:
                from server import conversion_progress
                conversion_progress[task_id] = {"progress": 20, "status": "Reading Excel file"}
            
            csv_path = os.path.join("converted", os.path.basename(xlsx_path).replace(".xlsx", "_converted.csv"))
            
            # Read Excel file in chunks if large
            try:
                df = pd.read_excel(xlsx_path, engine='openpyxl')
                
                if task_id:
                    conversion_progress[task_id] = {"progress": 70, "status": "Converting to CSV"}
                
                df.to_csv(csv_path, index=False)
                
                if task_id:
                    conversion_progress[task_id] = {"progress": 100, "status": "CSV conversion complete"}
                
                return csv_path
                
            except Exception as e:
                logger.error(f"Excel to CSV conversion failed: {e}")
                raise
                
        except Exception as e:
            logger.error(f"Excel to CSV conversion failed: {e}")
            raise

# Global optimized converter instance
optimized_converter = OptimizedConverter()

def get_optimized_converter() -> OptimizedConverter:
    """Get the global optimized converter instance"""
    return optimized_converter